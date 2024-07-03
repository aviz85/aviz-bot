import os
import io
import requests
import json
import faiss
import numpy as np
import pickle
from flask import current_app
from typing import Union, List, Dict, Any
from cohere import Client
from docx import Document
from PyPDF2 import PdfReader
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.core.schema import MetadataMode
from dotenv import load_dotenv

load_dotenv()

class VectorDB:
    VALID_EXTENSIONS = ('.docx', '.txt', '.pdf')
    CHUNK_SIZE = 100
    CHUNK_OVERLAP = 20
    COHERE_EMBED_MODEL = "embed-multilingual-v3.0"
    COHERE_RERANK_MODEL = "rerank-multilingual-v3.0"
    INITIAL_SEARCH_K = 100
    RERANK_TOP_N = 5

    def __init__(self, file_paths_or_urls: Union[str, List[str]]):
        print("Initializing VectorDB...")
        self.cohere_client = self._initialize_cohere()
        self.upload_folder = current_app.config['UPLOAD_FOLDER']
        self.manifest_file = os.path.join(self.upload_folder, "file_manifest.json")
        self.file_manifest = self._load_manifest()
        
        self.chunks = []
        self.embeddings = None
        self.index = None
        
        self._process_files(file_paths_or_urls)
        print("VectorDB initialization complete")

    def _initialize_cohere(self):
        cohere_api_key = os.getenv('COHERE_API_KEY')
        if not cohere_api_key:
            raise ValueError("COHERE_API_KEY not found in environment variables")
        return Client(cohere_api_key)

    def _load_manifest(self):
        if os.path.exists(self.manifest_file):
            with open(self.manifest_file, 'r') as f:
                return json.load(f)
        return {}

    def _save_manifest(self):
        with open(self.manifest_file, 'w') as f:
            json.dump(self.file_manifest, f)

    def _process_files(self, file_paths_or_urls):
        files = file_paths_or_urls if isinstance(file_paths_or_urls, list) else [file_paths_or_urls]
        for file_path in files:
            if self._is_valid_file(file_path):
                base_name = self._get_base_name(file_path)
                if self._should_process_file(file_path, base_name):
                    print(f"Processing file: {file_path}")
                    self._process_single_file(file_path, base_name)
                else:
                    print(f"Loading existing data for: {file_path}")
                    self._load_existing_data(base_name)
            else:
                print(f"Skipping invalid file: {file_path}")
        
        self._combine_data()

    def _is_valid_file(self, file_path):
        return file_path.lower().endswith(self.VALID_EXTENSIONS)

    def _get_base_name(self, file_path):
        return os.path.splitext(os.path.basename(file_path))[0]

    def _should_process_file(self, file_path, base_name):
        if base_name not in self.file_manifest:
            return True
        mod_time = os.path.getmtime(file_path)
        return self.file_manifest[base_name]['mod_time'] < mod_time

    def _process_single_file(self, file_path, base_name):
        content = self._load_files(file_path)
        chunks = self._split_text(content)
        embeddings = self._create_embeddings(chunks)
        
        chunks_file = os.path.join(self.upload_folder, f"{base_name}_chunks.pkl")
        embeddings_file = os.path.join(self.upload_folder, f"{base_name}_embeddings.npy")
        
        with open(chunks_file, 'wb') as f:
            pickle.dump(chunks, f)
        np.save(embeddings_file, embeddings)
        
        self.file_manifest[base_name] = {
            'mod_time': os.path.getmtime(file_path),
            'chunks_file': chunks_file,
            'embeddings_file': embeddings_file,
            'original_file': file_path
        }
        self._save_manifest()

    def _load_files(self, file_path_or_url: str) -> str:
        print(f"Loading file from: {file_path_or_url}")
        if file_path_or_url.startswith(('http://', 'https://')):
            response = requests.get(file_path_or_url)
            file_content = response.content
            return self._process_file_content(file_content, file_path_or_url)
        else:
            file_extension = os.path.splitext(file_path_or_url)[1].lower()
            if file_extension == '.docx':
                return self._process_docx(file_path_or_url)
            elif file_extension == '.pdf':
                return self._process_pdf(file_path_or_url)
            elif file_extension == '.txt':
                return self._process_txt(file_path_or_url)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

    def _process_docx(self, file_path: str) -> str:
        print(f"Processing DOCX file: {file_path}")
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

    def _process_pdf(self, file_path: str) -> str:
        print(f"Processing PDF file: {file_path}")
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            return "\n".join([page.extract_text() for page in pdf_reader.pages])

    def _process_txt(self, file_path: str) -> str:
        print(f"Processing TXT file: {file_path}")
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def _process_file_content(self, content: bytes, file_name: str) -> str:
        file_extension = os.path.splitext(file_name)[1].lower()
        if file_extension == '.docx':
            return self._process_docx_content(content)
        elif file_extension == '.pdf':
            return self._process_pdf_content(content)
        elif file_extension == '.txt':
            return content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def _process_docx_content(self, content: bytes) -> str:
        doc = Document(io.BytesIO(content))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    def _process_pdf_content(self, content: bytes) -> str:
        pdf_reader = PdfReader(io.BytesIO(content))
        return "\n".join([page.extract_text() for page in pdf_reader.pages])

    def _load_existing_data(self, base_name):
        file_info = self.file_manifest[base_name]
        with open(file_info['chunks_file'], 'rb') as f:
            chunks = pickle.load(f)
        embeddings = np.load(file_info['embeddings_file'])
        return chunks, embeddings

    def _combine_data(self):
        all_chunks = []
        all_embeddings = []
        for base_name, file_info in self.file_manifest.items():
            chunks, embeddings = self._load_existing_data(base_name)
            all_chunks.extend(chunks)
            all_embeddings.append(embeddings)
        
        self.chunks = all_chunks
        self.embeddings = np.vstack(all_embeddings) if all_embeddings else None
        self.index = self._create_faiss_index(self.embeddings)

    def _split_text(self, content: str) -> List[str]:
        document = LlamaDocument(text=content)
        parser = SimpleNodeParser.from_defaults(chunk_size=self.CHUNK_SIZE, chunk_overlap=self.CHUNK_OVERLAP)
        nodes = parser.get_nodes_from_documents([document])
        return [node.get_content(metadata_mode=MetadataMode.EMBED) for node in nodes]

    def _create_embeddings(self, chunks: List[str]) -> np.ndarray:
        print("Creating embeddings")
        embeddings = self.cohere_client.embed(
            texts=chunks,
            model=self.COHERE_EMBED_MODEL,
            input_type="search_document"
        ).embeddings
        return np.array(embeddings, dtype=np.float32)

    def _create_faiss_index(self, embeddings: np.ndarray) -> faiss.IndexFlatL2:
        print("Creating FAISS index")
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings)
        return index

    def search(self, queries: Union[str, List[str]], do_rerank: bool = True) -> List[Dict[str, Any]]:
        if isinstance(queries, str):
            queries = [queries]
        
        print(f"Searching for queries: {queries}")
        
        query_embeddings = self.cohere_client.embed(
            texts=queries,
            model=self.COHERE_EMBED_MODEL,
            input_type="search_query"
        ).embeddings
        
        all_indices = []
        all_distances = []
        for query_embedding in query_embeddings:
            distances, indices = self.index.search(np.array([query_embedding]), self.INITIAL_SEARCH_K)
            all_indices.extend(indices[0])
            all_distances.extend(distances[0])
        
        unique_indices = list(dict.fromkeys(all_indices))
        initial_results = [self.chunks[i] for i in unique_indices[:self.INITIAL_SEARCH_K]]
        
        if not do_rerank:
            return [{"text": initial_results[i], "relevance_score": all_distances[i], "original_index": int(unique_indices[i])} for i in range(min(len(initial_results), self.RERANK_TOP_N))]
        
        reranked_results = self._rerank(queries, initial_results)
        
        final_results = []
        for result in reranked_results:
            final_results.append({
                "text": result['text'],
                "relevance_score": result['relevance_score'],
                "original_index": int(unique_indices[result['index']])
            })
        
        return final_results

    def _rerank(self, queries: List[str], initial_results: List[str]) -> List[Dict[str, Any]]:
        all_reranked_results = []
        for query in queries:
            rerank_results = self.cohere_client.rerank(
                model=self.COHERE_RERANK_MODEL,
                query=query,
                documents=initial_results,
                top_n=self.RERANK_TOP_N,
                return_documents=False
            )
            
            for result in rerank_results.results:
                all_reranked_results.append({
                    'text': initial_results[result.index],
                    'index': result.index,
                    'relevance_score': result.relevance_score
                })
        
        unique_results = {result['text']: result for result in all_reranked_results}
        sorted_results = sorted(unique_results.values(), key=lambda x: x['relevance_score'], reverse=True)
        
        return sorted_results[:self.RERANK_TOP_N]

if __name__ == "__main__":
    print("Starting VectorDB example")
    db = VectorDB(["path/to/your/file1.txt", "path/to/your/file2.docx"])
    results = db.search(["Query 1 here", "Query 2 here", "Query 3 here"], do_rerank=True)
    print(results)